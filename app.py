from fastapi import FastAPI, UploadFile, File, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
import os
import uuid
import shutil
import time
import gc
import base64
import json
import asyncio
from datetime import datetime
from langchain_core.documents import Document

from langchain_mistralai import ChatMistralAI
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface.embeddings import HuggingFaceEndpointEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
import fitz  # PyMuPDF

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://vibieprince.github.io/"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Core Processing Models
embedding_model = HuggingFaceEndpointEmbeddings(model="sentence-transformers/all-MiniLM-L6-v2")
llm = ChatMistralAI(model="mistral-small-latest", temperature=0)

# Multi-Model OCR Orchestration Layers
# asyncio.wait_for(timeout=90) in the extraction loop is the primary hard deadline.
# ChatMistralAI additionally accepts a native `timeout` (seconds) that caps its
# internal HTTP client, cutting off LangChain retries earlier at the transport level.
# ChatGoogleGenerativeAI has no recognized top-level timeout field; the asyncio
# deadline covers it completely.
gemini_ocr_llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)
mistral_ocr_llm = ChatMistralAI(model="pixtral-12b-2409", temperature=0, timeout=45)

# Environment Safe Threshold Parameters (Applied to OCR ONLY)
MAX_FILE_SIZE_MB = 10
MAX_PAGE_LIMIT = 10

# Global thread-safe state broadcaster for frontend SSE tracking
processing_states = {}

prompt = ChatPromptTemplate.from_messages([
    ("system", (
        "You are a strict study assistant. Use ONLY the provided context to answer questions. "
        "If the answer is not contained within the context, specifically state: "
        "'I am sorry, but the uploaded document does not contain information regarding this.' "
        "Do not use outside knowledge or mention other companies/topics not in the text. "
        "Use markdown for formatting."
        "However if someone asks you general questions like how are you, what do you do, what's the day today, what will be the day on the date mentioned in the document, etc... you have to answer accordingly."
    )),
    ("human", "Context:\n{context}\n\nQuestion: {question}")
])

def update_state(session_id: str, status: str, progress: int, message: str, is_error: bool = False):
    processing_states[session_id] = {
        "status": status,
        "progress": progress,
        "message": message,
        "is_error": is_error,
        "timestamp": time.time()
    }

async def scheduled_cleanup(session_id: str):
    """Async cleanup task — runs after a 10-minute grace period without blocking the event loop."""
    print(f"⏳ Cleanup timer started for session: {session_id} (10 minutes)")
    await asyncio.sleep(600)

    db_path = f"db/{session_id}"
    temp_file = f"temp/{session_id}.pdf"

    gc.collect()

    if os.path.exists(db_path):
        for i in range(5):
            try:
                shutil.rmtree(db_path)
                print(f"✅ SUCCESS: Deleted DB folder for {session_id}")
                break
            except PermissionError:
                print(f"⚠️ Retrying DB deletion for {session_id}...")
                await asyncio.sleep(2)
            except Exception as e:
                print(f"❌ ERROR deleting DB: {e}")
                break

    if os.path.exists(temp_file):
        try:
            os.remove(temp_file)
            print(f"✅ SUCCESS: Deleted PDF file for {session_id}")
        except Exception as e:
            print(f"❌ ERROR deleting PDF: {e}")

    if session_id in processing_states:
        del processing_states[session_id]

    print(f"🏁 Finished cleanup routine for {session_id}")


async def run_pipeline(session_id: str, upload_path: str, db_path: str, filename: str, filename_lower: str):
    """
    Heavy async pipeline: OCR extraction → text splitting → Chroma vectorization.
    Runs entirely as a BackgroundTask so /upload can return the session_id immediately.
    """
    try:
        update_state(session_id, "starting", 5, "Starting initialization protocols...")

        cleaned_docs = []

        if filename_lower.endswith(('.doc', '.docx')):
            # ----------------------------------------------------------------
            # Structural Parsing Pass For DOCX Native Strings (python-docx)
            # ----------------------------------------------------------------
            update_state(session_id, "parsing", 20, "Parsing DOCX structural layout...")
            try:
                import docx as _docx
                doc_struct = _docx.Document(upload_path)
                # Filter out empty/whitespace-only paragraphs to avoid indexing blank lines
                combined_text = "\n".join([p.text for p in doc_struct.paragraphs if p.text.strip()])
                if not combined_text.strip():
                    update_state(session_id, "failed", 0, "The uploaded document appears to contain no readable text. Please check the file contents.", is_error=True)
                    if os.path.exists(upload_path): os.remove(upload_path)
                    return
                cleaned_docs.append(Document(page_content=combined_text, metadata={"source": filename, "page": 1}))
            except Exception as e:
                # Do NOT fall back to raw binary reads — a .docx is a ZIP archive and its
                # raw bytes are XML/PK headers, not human-readable text.
                update_state(session_id, "failed", 0, f"Failed to parse DOCX structure. Ensure the file is a valid Word document. Detail: {e}", is_error=True)
                if os.path.exists(upload_path): os.remove(upload_path)
                return
        else:
            # ----------------------------------------------------------------
            # Standard Native PDF Engine Processing Path
            # ----------------------------------------------------------------
            try:
                doc = fitz.open(upload_path)
            except Exception as e:
                update_state(session_id, "failed", 0, f"Failed to read PDF structure: {e}", is_error=True)
                if os.path.exists(upload_path): os.remove(upload_path)
                return

            num_pages = len(doc)

            # ----------------------------------------------------------------
            # ⚡ FAST-SCAN: Determine if document actually needs OCR
            # ----------------------------------------------------------------
            requires_ocr = False
            native_texts = []

            for i in range(num_pages):
                text = " ".join(doc[i].get_text().split()).strip()
                native_texts.append(text)
                if len(text) < 40:
                    requires_ocr = True

            active_ocr_engine = "gemini"

            # ----------------------------------------------------------------
            # ⚖️ CONDITIONAL LOGIC: Only restrict/ping if it's an Image PDF
            # ----------------------------------------------------------------
            if requires_ocr:
                update_state(session_id, "ocr_preflight", 15, "Scanned layout matrix detected. Activating OCR Preflight Diagnostics Engine...")
                print("🔍 Scanned images detected. OCR AI routing engaged.")

                # Immediate Pre-flight Check (Fallback routed upfront)
                try:
                    await asyncio.get_event_loop().run_in_executor(None, lambda: gemini_ocr_llm.invoke("ping"))
                    update_state(session_id, "ocr_processing", 25, "Primary Vision Core (Gemini) passed status diagnostics. Executing extraction blocks...")
                    print("✅ Primary AI Engine (Gemini) pre-flight passed.")
                except Exception as e:
                    if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e) or "Quota" in str(e):
                        print("⚠️ Gemini limit reached upfront. Instantly routing to Mistral Backup Engine...")
                        active_ocr_engine = "mistral"
                        update_state(session_id, "fallback_active", 30, "Gemini Exhaustion Rule Triggered (429 Rate Limit). Instantly routing workflows to Mistral-OCR Cluster Core...")
                        try:
                            await asyncio.get_event_loop().run_in_executor(None, lambda: mistral_ocr_llm.invoke("ping"))
                            print("✅ Backup Engine (Mistral Pixtral) pre-flight passed.")
                        except Exception as mistral_err:
                            update_state(session_id, "failed", 0, f"OCR Critical Fault: Both extraction engines are unresponsive. Mistral Log: {mistral_err}", is_error=True)
                            doc.close()
                            if os.path.exists(upload_path): os.remove(upload_path)
                            return
                    else:
                        update_state(session_id, "failed", 0, f"Primary Core Allocation Error: {str(e)}", is_error=True)
                        doc.close()
                        if os.path.exists(upload_path): os.remove(upload_path)
                        return
            else:
                update_state(session_id, "parsing", 40, "Standard layout configuration confirmed. Running recursive context parsing matrix layers...")
                print("📄 Standard text PDF detected. Bypassing size/page limits and AI pre-flights.")

            # ----------------------------------------------------------------
            # ⚙️ EXTRACTION LOOP
            # Per-page hard deadline: asyncio.wait_for kills the coroutine after
            # OCR_PAGE_TIMEOUT seconds even if LangChain's internal retry is still
            # running, guaranteeing the frontend is never frozen indefinitely.
            # ----------------------------------------------------------------
            OCR_PAGE_TIMEOUT = 90  # seconds per page across all retry attempts

            for i in range(num_pages):
                text = native_texts[i]

                if len(text) < 40:
                    pix = doc[i].get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("png")
                    img_base64 = base64.b64encode(img_bytes).decode("utf-8")

                    human_payload = HumanMessage(
                        content=[
                            {"type": "text", "text": "Extract all readable text from this document page. Return only extracted text exactly as written without explanations."},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                        ]
                    )

                    pct = int(40 + ((i + 1) / num_pages) * 35)
                    ocr_text = ""
                    page_failed = False

                    await asyncio.sleep(0.5)

                    # --- PRIMARY ENGINE ATTEMPT ---
                    try:
                        if active_ocr_engine == "gemini":
                            update_state(session_id, "ocr_processing", pct, f"Processing page layer {i+1} of {num_pages} via Gemini Multimodal Vision matrix...")
                            print(f"📷 Page {i+1}: Extracting with Gemini...")
                            response = await asyncio.wait_for(
                                asyncio.get_event_loop().run_in_executor(None, lambda: gemini_ocr_llm.invoke([human_payload])),
                                timeout=OCR_PAGE_TIMEOUT
                            )
                            ocr_text = response.content.strip()
                        else:
                            update_state(session_id, "fallback_active", pct, f"Processing page layer {i+1} of {num_pages} via Mistral Vision Backup cluster...")
                            print(f"📷 Page {i+1}: Extracting with Mistral Vision Backup...")
                            response = await asyncio.wait_for(
                                asyncio.get_event_loop().run_in_executor(None, lambda: mistral_ocr_llm.invoke([human_payload])),
                                timeout=OCR_PAGE_TIMEOUT
                            )
                            ocr_text = response.content.strip()

                    except (asyncio.TimeoutError, Exception) as e:
                        error_msg = str(e)
                        is_timeout = isinstance(e, asyncio.TimeoutError)
                        print(f"❌ OCR {'timeout' if is_timeout else 'error'} on page {i+1}: {error_msg}")

                        # Rate-limit on Gemini → try cascading to Mistral
                        if active_ocr_engine == "gemini" and not is_timeout and \
                                ("429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg or "Quota" in error_msg):
                            print(f"🔄 Rate limit hit on page {i+1}. Cascading to Mistral...")
                            active_ocr_engine = "mistral"
                            update_state(session_id, "fallback_active", pct,
                                         f"Cascade Error on page {i+1}! Shifting queue to Mistral Core...")
                            # --- MISTRAL CASCADE ATTEMPT ---
                            try:
                                response = await asyncio.wait_for(
                                    asyncio.get_event_loop().run_in_executor(None, lambda: mistral_ocr_llm.invoke([human_payload])),
                                    timeout=OCR_PAGE_TIMEOUT
                                )
                                ocr_text = response.content.strip()
                            except (asyncio.TimeoutError, Exception) as fallback_err:
                                fb_msg = "Request timed out" if isinstance(fallback_err, asyncio.TimeoutError) else str(fallback_err)
                                print(f"❌ Mistral cascade also failed on page {i+1}: {fb_msg}")
                                page_failed = True
                                update_state(session_id, "failed", 0,
                                             f"Network Pipeline Timeout: Both OCR engines failed on page {i+1}. "
                                             f"Mistral error — {fb_msg}", is_error=True)
                        else:
                            # Timeout or non-rate-limit error on whichever engine is active
                            label = "Request timed out after 90s" if is_timeout else error_msg
                            engine_name = "Gemini" if active_ocr_engine == "gemini" else "Mistral"
                            page_failed = True
                            update_state(session_id, "failed", 0,
                                         f"Network Pipeline Timeout: {engine_name} OCR dropped on page {i+1}. "
                                         f"Detail — {label}", is_error=True)

                    # If this page hit a terminal failure, abort the entire pipeline immediately
                    # so the SSE error is delivered and the frontend unlocks.
                    if page_failed:
                        doc.close()
                        if os.path.exists(upload_path):
                            try: os.remove(upload_path)
                            except: pass
                        return

                    if ocr_text:
                        cleaned_docs.append(Document(page_content=ocr_text, metadata={"source": filename, "page": i+1}))
                else:
                    cleaned_docs.append(Document(page_content=text, metadata={"source": filename, "page": i+1}))

            doc.close()

        # ----------------------------------------------------------------
        # ✅ VALIDATION & INDEXING
        # ----------------------------------------------------------------
        if not cleaned_docs or not cleaned_docs[0].page_content.strip():
            update_state(session_id, "failed", 0, "Parsing Pipeline Timeout: No readable character strings or image blocks could be synthesized.", is_error=True)
            if os.path.exists(upload_path): os.remove(upload_path)
            return

        update_state(session_id, "indexing", 85, "Structuring context chunk distributions and vector maps split rules...")
        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
        chunks = splitter.split_documents(cleaned_docs)

        # Run blocking Chroma operation in a thread executor so it doesn't block the event loop
        def _build_vectorstore():
            vs = Chroma.from_documents(
                documents=chunks,
                embedding=embedding_model,
                persist_directory=db_path
            )
            del vs

        await asyncio.get_event_loop().run_in_executor(None, _build_vectorstore)

        update_state(session_id, "completed", 100, "Analysis complete! Content blocks processed into localized vector store indexes.")
        print(f"✅ Pipeline completed for session: {session_id}")

        # Schedule async cleanup without blocking
        asyncio.create_task(scheduled_cleanup(session_id))

    except Exception as e:
        update_state(session_id, "failed", 0, f"Unexpected pipeline fault: {str(e)}", is_error=True)
        print(f"❌ Unhandled pipeline exception for {session_id}: {e}")
        if os.path.exists(upload_path): 
            try: os.remove(upload_path)
            except: pass


@app.get("/progress/{session_id}")
async def get_progress_stream(session_id: str):
    """Server-Sent Events (SSE) Endpoint delivering native pipeline status logs to frontend UI."""
    async def event_generator():
        while True:
            if session_id in processing_states:
                state = processing_states[session_id]
                # Construct clean payload dictionary
                payload = {
                    "status": state["status"],
                    "progress": state["progress"],
                    "message": state["message"],
                    "is_error": state["is_error"]
                }
                # Yield as a strict valid JSON string
                yield f"data: {json.dumps(payload)}\n\n"

                if state['progress'] == 100 or state['is_error']:
                    break
            else:
                fallback_payload = {
                    "status": "initializing",
                    "progress": 0,
                    "message": "Waiting for stream connection handshake...",
                    "is_error": False
                }
                yield f"data: {json.dumps(fallback_payload)}\n\n"
            await asyncio.sleep(0.5)

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.post("/upload")
async def upload_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    session_id = str(uuid.uuid4())

    # ----------------------------------------------------------------
    # 🛑 CRITICAL STAGE 0: IMMEDIATE FAIL-FAST PRE-FLIGHT VALIDATIONS
    # These checks happen synchronously and return errors instantly.
    # ----------------------------------------------------------------
    filename_lower = file.filename.lower()
    allowed_extensions = ('.pdf', '.doc', '.docx')
    if not filename_lower.endswith(allowed_extensions):
        return {
            "error": "Unsupported File Type! Access allowed to PDF and DOC/DOCX files only.",
            "session_id": session_id
        }

    # Pre-read stream for size verification
    try:
        file_bytes = await file.read()
    except Exception as e:
        return {
            "error": f"Streaming error reading input parameters: {str(e)}",
            "session_id": session_id
        }

    file_size_mb = len(file_bytes) / (1024 * 1024)

    # Enforce static allocation maximums early
    if file_size_mb > MAX_FILE_SIZE_MB:
        return {
            "error": f"File payload validation failure! Current file size ({file_size_mb:.2f}MB) exceeds system limitations of {MAX_FILE_SIZE_MB}MB.",
            "session_id": session_id
        }

    # Page count pre-flight for PDFs only
    if filename_lower.endswith('.pdf'):
        try:
            import io
            doc_check = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
            num_pages_check = len(doc_check)
            doc_check.close()
            if num_pages_check > MAX_PAGE_LIMIT:
                return {
                    "error": f"Document layout boundary rejection! File length ({num_pages_check} pages) exceeds current system limit rules ({MAX_PAGE_LIMIT} pages max).",
                    "session_id": session_id
                }
        except Exception as e:
            return {
                "error": f"Failed to read PDF structure during pre-flight: {e}",
                "session_id": session_id
            }

    # ----------------------------------------------------------------
    # ✅ PRE-FLIGHT PASSED — save file and hand off to background task
    # ----------------------------------------------------------------
    upload_path = f"temp/{session_id}.pdf"
    db_path = f"db/{session_id}"

    os.makedirs("temp", exist_ok=True)
    os.makedirs("db", exist_ok=True)

    with open(upload_path, "wb") as f:
        f.write(file_bytes)

    # Seed the state so the SSE stream has something to read immediately
    update_state(session_id, "starting", 1, "Starting initialization protocols...")

    # Offload the entire heavy pipeline — /upload returns immediately after this line
    # FastAPI BackgroundTasks natively supports async coroutine functions
    background_tasks.add_task(run_pipeline, session_id, upload_path, db_path, file.filename, filename_lower)

    # Return session_id instantly so the frontend can open the EventSource NOW
    return {"session_id": session_id}


@app.post("/chat")
async def chat(data: dict):
    query = data.get("query")
    session_id = data.get("session_id")
    db_path = f"db/{session_id}"

    if not session_id or not os.path.exists(db_path):
        return {"answer": "Session expired or invalid. Please re-upload your document."}

    vectorstore = Chroma(
        persist_directory=db_path,
        embedding_function=embedding_model
    )

    retriever = vectorstore.as_retriever(
        search_type="mmr", 
        search_kwargs={"k": 3, "fetch_k": 10}
    )

    docs = retriever.invoke(query)

    if not docs:
        return {"answer": "I am sorry, but I couldn't find any relevant sections in the document to answer that."}

    current_time_str = datetime.now().strftime("%A, %B %d, %Y")
    context_elements = [f"Current Live System Date/Time: {current_time_str}"]
    context_elements.extend(doc.page_content for doc in docs)
    context = "\n\n".join(context_elements)

    final_prompt = prompt.invoke({"context": context, "question": query})
    response = llm.invoke(final_prompt)

    del vectorstore
    gc.collect()

    return {"answer": response.content}


@app.get("/storage-status")
async def get_storage_status():
    temp_files = os.listdir("temp") if os.path.exists("temp") else []
    db_folders = os.listdir("db") if os.path.exists("db") else []

    return {
        "active_pdfs_count": len(temp_files),
        "active_pdfs": temp_files,
        "active_db_sessions_count": len(db_folders),
        "active_db_sessions": db_folders,
        "location_info": {
            "pdf_storage": os.path.abspath("temp"),
            "vector_storage": os.path.abspath("db")
        }
    }